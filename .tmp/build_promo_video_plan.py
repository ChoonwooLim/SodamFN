from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\WORK\SodamFN")
OUT = ROOT / "docs" / "marketing" / "SEMHANA_홍보영상_제작기획서.docx"
ICON = ROOT / "SodamApp" / "frontend" / "public" / "icons" / "icon-512x512.png"

NAVY = "0F172A"
SLATE = "334155"
MUTED = "64748B"
BLUE = "3B82F6"
TEAL = "06A884"
GREEN = "10B981"
AMBER = "F59E0B"
RED = "DC2626"
INK = "111827"
LIGHT = "F1F5F9"
PALE_BLUE = "EFF6FF"
PALE_TEAL = "ECFDF5"
WHITE = "FFFFFF"
BORDER = "CBD5E1"
FONT = "Malgun Gothic"


def set_run_font(run, size=11, bold=False, color=INK, italic=False, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6", inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in edges:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_table_geometry(table, widths, indent=120):
    assert sum(widths) == 9360, widths
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def keep_together(paragraph):
    paragraph.paragraph_format.keep_together = True


def set_para_border_left(paragraph, color=BLUE, size="18", space="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_text(doc, text, size=11, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8, line=1.333,
             keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if keep:
        keep_together(p)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_rich_text(doc, parts, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8,
                  line=1.333, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if keep:
        keep_together(p)
    for item in parts:
        r = p.add_run(item["text"])
        set_run_font(
            r,
            size=item.get("size", 11),
            bold=item.get("bold", False),
            color=item.get("color", INK),
            italic=item.get("italic", False),
        )
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = False
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=16, bold=True, color=NAVY)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=13, bold=True, color=BLUE)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True, color=SLATE)
    return p


def add_bullet(doc, text, num_id, level=0, color=INK, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color)
    return p


def add_number(doc, text, num_id):
    return add_bullet(doc, text, num_id)


def ensure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids or [0]) + 1
    next_num = max(num_ids or [0]) + 1

    def create(fmt, marker, abs_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        fonts.set(qn("w:eastAsia"), FONT)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)

    create("bullet", "●", next_abs, next_num)
    bullet_num = next_num
    create("decimal", "%1.", next_abs + 1, next_num + 1)
    decimal_num = next_num + 1
    return bullet_num, decimal_num


def add_callout(doc, label, title, body, fill=PALE_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_together = True
    set_para_border_left(p, color=accent, size="20", space="10")
    set_paragraph_shading(p, fill)
    r = p.add_run(label.upper() + "\n")
    set_run_font(r, size=9, bold=True, color=accent)
    r2 = p.add_run(title + "\n")
    set_run_font(r2, size=13, bold=True, color=NAVY)
    r3 = p.add_run(body)
    set_run_font(r3, size=10.5, color=SLATE)
    return p


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2400, 6960])
    set_table_borders(table, color=BORDER, size="5")
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        set_cell_shading(left, LIGHT)
        for p in left.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(label)
            set_run_font(r, size=10, bold=True, color=SLATE)
        for p in right.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(value)
            set_run_font(r, size=10, color=INK)
    add_text(doc, "", size=2, after=3)
    return table


def add_matrix(doc, headers, rows, widths, header_fill=NAVY, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table, color=BORDER, size="5")
    mark_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        if ridx % 2 == 1:
            for c in cells:
                set_cell_shading(c, "F8FAFC")
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, bold=(idx == 0), color=INK)
    add_text(doc, "", size=2, after=4)
    return table


def add_section_banner(doc, number, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    set_paragraph_shading(p, NAVY)
    r1 = p.add_run(f"  {number}  ")
    set_run_font(r1, size=10, bold=True, color=TEAL)
    r2 = p.add_run(label)
    set_run_font(r2, size=10, bold=True, color=WHITE)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for style_name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, SLATE, 8, 4),
]:
    st = styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
hr = hp.add_run("SEMHANA  ·  PROMOTIONAL FILM BLUEPRINT")
set_run_font(hr, size=8.5, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run("SEMHANA  |  2026.07.23    ")
set_run_font(fr, size=8, color=MUTED)
add_page_field(fp)

bullet_num, decimal_num = ensure_numbering(doc)

# Cover
add_text(doc, "PROMOTIONAL FILM BLUEPRINT", size=9.5, color=BLUE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=20, line=1.0)
add_text(doc, "SEM\nHANA", size=13, color=WHITE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=18, line=0.86)
cover_mark = doc.paragraphs[-1]
cover_mark.paragraph_format.left_indent = Inches(2.72)
cover_mark.paragraph_format.right_indent = Inches(2.72)
set_paragraph_shading(cover_mark, NAVY)

add_text(doc, "셈하나 홍보영상\n제작 기획서", size=29, color=NAVY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10, line=1.1)
add_text(doc, "“복잡한 가게 운영을, 하나로.”", size=16, color=BLUE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=16, line=1.0)
add_text(doc, "사장님의 하루에 여유를 돌려주는 60초 브랜드 필름", size=12.5, color=SLATE,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=34, line=1.2)

cover_meta = doc.add_table(rows=3, cols=2)
set_table_geometry(cover_meta, [4680, 4680], indent=120)
remove_table_borders(cover_meta)
cover_rows = [
    ("프로젝트", "셈하나 브랜드 인지도·체험 전환 영상"),
    ("핵심 타깃", "40~50대 소상공인·매장 운영자"),
    ("권장 구성", "60초 메인 + 30초·15초·6초 파생본"),
]
for i, (left, right) in enumerate(cover_rows):
    c1, c2 = cover_meta.rows[i].cells
    for c in (c1, c2):
        set_cell_margins(c, top=90, bottom=90, start=120, end=120)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(left)
    set_run_font(r1, size=9.5, bold=True, color=MUTED)
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(right)
    set_run_font(r2, size=9.5, color=INK)

add_text(doc, "Prepared  ·  2026.07.23", size=9, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=0, line=1.0)
doc.add_page_break()

# 1. Executive
add_section_banner(doc, "01", "EXECUTIVE BRIEF")
add_h1(doc, "한 줄로 말하면")
add_callout(
    doc,
    "Recommended Creative Direction",
    "영업이 끝나도 일이 끝나지 않는 사장님의 밤을, 셈하나가 바꾼다.",
    "기능을 나열하는 앱 광고가 아니라, ‘복잡함 → 자동 정리 → 안심 → 여유’의 감정 변화를 보여준다. "
    "시청자는 10초 안에 자신의 문제를 발견하고, 45초 안에 제품의 효용을 이해하며, 60초에 체험 행동을 선택하게 된다.",
)

add_h2(doc, "기획 핵심")
add_key_value_table(doc, [
    ("목표", "처음 보는 사람에게 ‘내 가게에도 꼭 필요하다’는 욕구를 만들고 무료 체험 또는 상담으로 연결"),
    ("주인공", "영업·정산·직원·자재를 동시에 챙겨야 하는 40~50대 사장님"),
    ("브랜드 감정", "신뢰감 · 심플함 · 따뜻함 — 숫자를 대신 처리해주는 든든한 비서"),
    ("핵심 약속", "흩어진 가게 운영 데이터를 한곳에 모아, 오늘의 상황을 이해하고 다음 행동까지 쉽게"),
    ("대표 장면", "영수증 촬영 → AI 분류 → 지출 반영 → 손익 변화가 한 흐름으로 이어지는 화면"),
    ("최종 CTA", "1순위 ‘무료로 시작하기’ / 가입 동선 미완성 시 ‘도입 상담 신청’로 교체"),
])

add_h2(doc, "성공하는 영상의 세 가지 조건")
add_number(doc, "첫 3초에 ‘내 얘기’여야 한다 — 계산기, 영수증 더미, 단체방 메시지, 늦은 밤을 빠르게 제시한다.", decimal_num)
add_number(doc, "기능은 반드시 결과와 붙인다 — 버튼이 아니라 ‘시간 절약·놓침 방지·한눈에 파악’을 보여준다.", decimal_num)
add_number(doc, "마지막은 안도감으로 끝낸다 — 사장님이 앱을 닫고 가게 불을 끄는 장면이 제품의 진짜 가치다.", decimal_num)

# 2. Strategy
add_section_banner(doc, "02", "AUDIENCE & MESSAGE STRATEGY")
add_h1(doc, "누구의 어떤 마음을 움직일 것인가")
add_h2(doc, "핵심 타깃 인사이트")
add_bullet(doc, "영업은 잘하지만 숫자 관리가 밀리면 ‘내가 가게를 제대로 보고 있나’ 하는 불안이 생긴다.", bullet_num)
add_bullet(doc, "복잡한 ERP보다 스마트폰에서 바로 이해되는 익숙한 화면을 원한다.", bullet_num)
add_bullet(doc, "기능의 개수보다 ‘오늘 얼마가 남았는지’, ‘급여가 맞는지’, ‘무엇을 주문해야 하는지’를 빠르게 알고 싶다.", bullet_num)
add_bullet(doc, "앱을 배우는 데 시간을 쓰고 싶지 않다. 보는 순간 사용법이 짐작되어야 한다.", bullet_num)

add_h2(doc, "메시지 구조")
add_matrix(
    doc,
    ["단계", "시청자 머릿속 질문", "영상의 답"],
    [
        ("공감", "“왜 정리는 늘 마감 후에도 끝나지 않지?”", "흩어진 영수증·매출·급여·재고의 현실을 짧고 정확하게 묘사"),
        ("발견", "“이걸 정말 한곳에서 할 수 있나?”", "실제 셈하나 화면을 손의 행동과 연결해 증명"),
        ("확신", "“내가 써도 어렵지 않을까?”", "큰 글씨, 한눈에 보이는 카드, 익숙한 모바일 조작을 강조"),
        ("욕구", "“나도 오늘부터 덜 힘들 수 있겠다.”", "사장님의 표정과 퇴근 장면으로 정서적 보상 제공"),
        ("행동", "“어디서 시작하지?”", "단 하나의 CTA와 짧은 URL·QR을 3초 이상 노출"),
    ],
    [1150, 3200, 5010],
)

add_h2(doc, "가치 증명 우선순위")
add_number(doc, "영수증 촬영과 AI 자동 분류 — 가장 즉각적이고 시각적인 ‘와우’ 장면", decimal_num)
add_number(doc, "매출·지출·손익 통합 — 사장님이 가장 알고 싶은 ‘그래서 얼마가 남았나’", decimal_num)
add_number(doc, "구매요청·자재·재고 흐름 — 매장 운영의 실무가 이어지는 차별점", decimal_num)
add_number(doc, "급여·출퇴근·직원앱 — 사장님과 직원 모두 편해지는 신뢰 장면", decimal_num)

add_h3(doc, "카피 가드레일")
add_bullet(doc, "검증되지 않은 절대 표현(‘100% 정확’, ‘완전 자동’, ‘모든 업종’)은 사용하지 않는다.", bullet_num)
add_bullet(doc, "실제 제공 중인 기능만 촬영하고, 연동 상태에 따라 수동 단계가 있는 기능은 화면에서 오해가 없게 한다.", bullet_num)
add_bullet(doc, "세무·노무 결과를 보장하는 표현보다 ‘계산과 정리를 돕는다’는 제품 역할을 명확히 한다.", bullet_num)

# 3. Creative concept
add_section_banner(doc, "03", "CREATIVE CONCEPT")
add_h1(doc, "크리에이티브 콘셉트: “사장님의 밤을 돌려드립니다”")
add_text(
    doc,
    "폐점 후에도 계산기와 영수증 앞에 남아 있는 사장님. 카메라는 복잡함을 짧고 촘촘하게 보여준 뒤, "
    "셈하나를 여는 순간부터 화면과 호흡이 정돈된다. 마지막에는 숫자가 아니라 ‘제시간에 가게 문을 닫는 사람’을 보여준다.",
)

add_h2(doc, "감정 곡선")
add_matrix(
    doc,
    ["0–8초", "8–22초", "22–45초", "45–60초"],
    [
        ("압박·공감", "발견·놀라움", "이해·확신", "안도·행동"),
        ("빠른 컷, 소음, 어두운 톤", "촬영 한 번, UI가 정리되는 리듬", "핵심 기능 3개를 결과 중심으로", "따뜻한 매장 빛, 단일 CTA"),
    ],
    [2340, 2340, 2340, 2340],
    header_fill=SLATE,
    font_size=9.2,
)

add_h2(doc, "비주얼 문법")
add_bullet(doc, "현실 장면: 따뜻한 다큐멘터리 톤, 실제 가게의 손·표정·영업 동선을 사용한다.", bullet_num)
add_bullet(doc, "제품 장면: 스마트폰 세로 화면을 115% 이상 확대하고, 손가락 탭과 데이터 변화를 한 장면에 연결한다.", bullet_num)
add_bullet(doc, "색 변화: 문제 구간은 Slate 900 중심, 해결 구간은 Blue 500과 Teal 포인트로 밝아지게 한다.", bullet_num)
add_bullet(doc, "모션: 카드가 부드럽게 정렬되고 숫자가 자연스럽게 갱신되는 절제된 움직임. 과한 3D·바운스 효과는 피한다.", bullet_num)
add_bullet(doc, "자막: 핵심어 3~6단어, 최대 2줄. 1080p 세로 영상 기준 44px 이상과 충분한 대비를 확보한다.", bullet_num)

add_h2(doc, "사운드 방향")
add_bullet(doc, "도입: 계산기 버튼, 종이 넘김, 알림음이 겹치는 현실 소리로 압박감을 만든다.", bullet_num)
add_bullet(doc, "전환: 영수증 촬영 셔터와 함께 소음이 정리되고, 92–104 BPM의 따뜻한 미니멀 비트가 시작된다.", bullet_num)
add_bullet(doc, "엔딩: 음악을 한 박자 비우고 브랜드 문장과 로고 사운드를 선명하게 남긴다.", bullet_num)

# 4. Storyboard
add_section_banner(doc, "04", "60-SECOND HERO STORYBOARD")
add_h1(doc, "60초 메인 영상 구성안")
story_rows = [
    ("0–3초", "늦은 밤. 계산기, 영수증 더미, 단체방 알림을 0.5초 컷으로 연결", "“장사는 끝났는데…”", "강제 공감"),
    ("3–8초", "사장님이 한숨을 쉬며 숫자를 맞춘다. 시계는 밤 11시 48분", "“사장님의 일은 왜 끝나지 않을까요?”", "문제 정의"),
    ("8–15초", "스마트폰으로 영수증 촬영. AI가 거래처·금액·항목을 읽고 정리", "“영수증은 찍기만 하세요.”", "첫 와우"),
    ("15–23초", "매출·지출 카드가 한 화면에 모이고 손익 숫자가 갱신", "“매출과 지출은 한곳에 모이고,”", "핵심 효용"),
    ("23–31초", "구매요청서 체크 → 거래처별 주문 → 재고 카드로 전환", "“주문과 재고도 자연스럽게 이어집니다.”", "운영 연결"),
    ("31–40초", "직원앱 출퇴근·급여명세서 확인. 사장님 화면에는 급여 정리 완료", "“직원은 자기 앱에서, 사장님은 한 화면에서.”", "양쪽 편의"),
    ("40–50초", "대시보드 핵심 숫자 확대. 사장님 표정이 편안해진다", "“오늘 번 돈과 남은 돈을 바로 봅니다.”", "확신"),
    ("50–56초", "가게 불을 끄고 문을 닫는다. 가족에게 ‘지금 출발’ 메시지", "“사장님의 밤을 돌려드릴게요.”", "정서적 보상"),
    ("56–60초", "앱 아이콘·브랜드명·CTA·QR. 화면은 단순하게 유지", "“복잡한 가게 운영을, 하나로. 셈하나.”", "행동 유도"),
]
add_matrix(
    doc,
    ["구간", "화면·액션", "내레이션/자막", "역할"],
    story_rows,
    [980, 3870, 3200, 1310],
    font_size=8.35,
)

add_h2(doc, "연출 포인트")
add_bullet(doc, "UI 녹화는 화면만 따로 보여주지 말고, 실제 손의 행동과 결과 화면을 매치 컷으로 연결한다.", bullet_num)
add_bullet(doc, "숫자 데이터는 데모 계정으로 준비하되 현실적인 금액과 거래 흐름을 사용한다.", bullet_num)
add_bullet(doc, "기능당 5~8초 이상 머물지 않는다. ‘동작 1개 + 결과 1개’만 남긴다.", bullet_num)
add_bullet(doc, "엔딩 CTA는 최소 3초 유지하고 URL·QR·버튼 문구를 하나로 통일한다.", bullet_num)

# 5. Script
add_section_banner(doc, "05", "VOICE-OVER & ON-SCREEN COPY")
add_h1(doc, "권장 내레이션 완성본")
script = [
    ("0–8초", "장사는 끝났는데, 사장님의 일은 왜 끝나지 않을까요?"),
    ("8–15초", "흩어진 영수증, 이제 찍기만 하세요. 셈하나가 읽고 정리를 돕습니다."),
    ("15–23초", "현금, 카드, 배달 매출과 지출은 한곳에 모이고,"),
    ("23–31초", "구매 요청과 재고, 거래처 관리까지 자연스럽게 이어집니다."),
    ("31–40초", "직원은 출퇴근과 급여명세서를 자기 앱에서 확인하고,"),
    ("40–50초", "사장님은 오늘 번 돈과 남은 돈을 바로 봅니다."),
    ("50–56초", "복잡한 숫자 대신, 사장님의 시간을 돌려드릴게요."),
    ("56–60초", "복잡한 가게 운영을, 하나로. 사장님의 숫자 비서, 셈하나."),
]
for timing, line in script:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    set_para_border_left(p, color=TEAL, size="14", space="8")
    r1 = p.add_run(f"{timing}  ")
    set_run_font(r1, size=9.5, bold=True, color=TEAL)
    r2 = p.add_run(line)
    set_run_font(r2, size=11.5, color=INK)

add_h2(doc, "화면 자막 키워드")
add_matrix(
    doc,
    ["장면", "메인 자막", "보조 자막"],
    [
        ("문제", "마감 후에도 끝나지 않는 일", "영수증 · 매출 · 급여 · 재고"),
        ("영수증", "찍으면 정리가 시작됩니다", "AI가 항목을 읽고 분류"),
        ("손익", "오늘 번 돈, 오늘 남은 돈", "매출과 지출을 한눈에"),
        ("자재", "요청부터 주문·재고까지", "매장 운영 흐름을 하나로"),
        ("직원", "직원도, 사장님도 더 편하게", "출퇴근 · 급여명세서 · 계약"),
        ("엔딩", "복잡한 가게 운영을, 하나로", "셈하나"),
    ],
    [1500, 3900, 3960],
)

add_h2(doc, "대체 슬로건 후보")
add_bullet(doc, "사장님의 숫자 비서, 셈하나.", bullet_num)
add_bullet(doc, "오늘의 장사를, 한눈에.", bullet_num)
add_bullet(doc, "장사는 사장님이. 복잡한 정리는 셈하나가.", bullet_num)
add_bullet(doc, "가게 운영의 복잡함을, 하나로.", bullet_num)
add_text(doc, "최종 추천: “복잡한 가게 운영을, 하나로.” — 브랜드명 ‘셈하나’와 가장 자연스럽게 결합되고 기능 확장에도 오래 사용할 수 있다.",
         size=10.5, color=SLATE, italic=True, before=2, after=8)

# 6. Cutdowns
add_section_banner(doc, "06", "CUTDOWNS & FORMAT SYSTEM")
add_h1(doc, "한 번 촬영해 네 가지 광고로 확장")
add_matrix(
    doc,
    ["버전", "목적", "구성", "권장 CTA"],
    [
        ("60초 16:9", "홈페이지·유튜브·영업 미팅", "문제 → 기능 4개 → 감정적 결말", "무료로 시작하기"),
        ("30초 9:16", "인스타 릴스·유튜브 쇼츠", "3초 훅 → 영수증 → 손익 → 직원앱 → CTA", "지금 체험하기"),
        ("15초 9:16", "성과형 광고·리타기팅", "영수증 촬영 전후 + 대시보드 결과", "영수증 한 장부터"),
        ("6초 9:16", "범퍼 광고·리마인드", "‘찍고 → 정리되고 → 보인다’ 3컷", "셈하나"),
    ],
    [1500, 2200, 4050, 1610],
)

add_h2(doc, "30초 숏폼 러닝오더")
add_number(doc, "0–3초: ‘이 영수증들, 오늘도 직접 정리하세요?’", decimal_num)
add_number(doc, "3–10초: 촬영 → AI 분류 → 지출 반영", decimal_num)
add_number(doc, "10–18초: 매출·지출·손익 한 화면", decimal_num)
add_number(doc, "18–25초: 직원앱과 구매요청·재고를 빠르게 연결", decimal_num)
add_number(doc, "25–30초: ‘복잡한 가게 운영을, 하나로. 셈하나.’ + CTA", decimal_num)

add_h2(doc, "첫 3초 훅 A/B 테스트")
add_matrix(
    doc,
    ["버전", "오프닝 카피", "추천 대상"],
    [
        ("A · 고통", "“마감했는데, 아직 퇴근 못 하셨죠?”", "브랜드 첫 노출·공감형"),
        ("B · 놀라움", "“영수증 한 장 찍었을 뿐인데…”", "기능 관심·성과형"),
        ("C · 결과", "“오늘 얼마 남았는지 바로 아세요?”", "손익 관심·사장님 커뮤니티"),
    ],
    [1500, 4750, 3110],
)

# 7. Capture plan
add_section_banner(doc, "07", "SHOT LIST & ASSET PLAN")
add_h1(doc, "필수 촬영·화면 자산")
add_h2(doc, "실사 촬영")
add_bullet(doc, "40~50대 실제 사장님 또는 신뢰감 있는 모델 1명, 20~30대 직원 1명", bullet_num)
add_bullet(doc, "영업 마감 직후의 매장, 카운터, 영수증, 계산기, 재고 선반, 출입문", bullet_num)
add_bullet(doc, "손 클로즈업: 영수증 촬영, 구매요청 체크, 대시보드 스크롤, 직원앱 출퇴근", bullet_num)
add_bullet(doc, "감정 컷: 피곤함 → 이해 → 안도 → 가게 문을 닫고 퇴근", bullet_num)

add_h2(doc, "필수 UI 화면")
add_matrix(
    doc,
    ["우선", "화면", "촬영 동작", "준비사항"],
    [
        ("1", "영수증 보관함", "촬영·업로드 → AI 분석 → 분류 완료", "데모 영수증 3종, 처리 속도 사전 확인"),
        ("2", "대시보드/손익", "매출·지출·이익 카드 확대", "현실적인 월간 데이터, 숫자 노출 검수"),
        ("3", "구매요청·자재", "품목 체크 → 수량 → 요청 완료", "단가 미노출 직원 화면 포함"),
        ("4", "재고·거래처", "재고 확인 → 발주 흐름", "품목 이미지와 거래처명 데모화"),
        ("5", "급여·직원앱", "출퇴근 → 급여명세서 확인", "실명·실급여 대신 가상 인물 사용"),
    ],
    [800, 2100, 3500, 2960],
    font_size=8.7,
)

add_h2(doc, "화면 녹화 규격")
add_bullet(doc, "관리자 앱: 1920×1080 브라우저 캡처와 1080×1920 모바일 캡처를 모두 확보", bullet_num)
add_bullet(doc, "직원앱: 실제 스마트폰 1080×1920, 알림·배터리·개인 계정 정보는 숨김", bullet_num)
add_bullet(doc, "터치 표시와 마우스 커서 스타일을 통일하고 60fps로 녹화", bullet_num)
add_bullet(doc, "편집용 안전영역: 세로 영상 상·하단 UI 영역을 피해 핵심 텍스트를 중앙 80%에 배치", bullet_num)

add_callout(
    doc,
    "Data Safety",
    "실데이터는 영상에 절대 사용하지 않는다.",
    "가상 매장·가상 직원·가상 급여·가상 계좌·가상 영수증으로 구성한 전용 데모 계정을 만든다. "
    "촬영 전 프레임 단위로 전화번호, 주소, 계좌, 주민등록 정보, 실제 매출·급여가 노출되지 않는지 확인한다.",
    fill="FEF2F2",
    accent=RED,
)

# 8. Production
add_section_banner(doc, "08", "PRODUCTION ROADMAP")
add_h1(doc, "권장 제작 방식: 실사 + 실제 UI + 절제된 모션그래픽")
add_text(
    doc,
    "셈하나의 신뢰를 가장 빠르게 만드는 조합이다. 실사는 감정과 현실감을, 실제 UI는 제품의 진정성을, "
    "모션그래픽은 복잡한 데이터 흐름의 이해를 담당한다. AI 생성 인물이나 과장된 3D 화면은 메인 영상에서 보조적으로만 사용한다.",
)

add_h2(doc, "6단계 제작 일정")
add_matrix(
    doc,
    ["단계", "기간", "핵심 작업", "완료 기준"],
    [
        ("1. 확정", "1일", "목표·CTA·슬로건·기능 우선순위 승인", "한 문장 메시지와 전환 경로 확정"),
        ("2. 준비", "1–2일", "데모 계정·데이터·소품·출연·장소 준비", "모든 UI 동선 리허설 완료"),
        ("3. 촬영", "1일", "실사 A-roll/B-roll, 폰 화면, 현장음", "60초·숏폼에 필요한 세로·가로 컷 확보"),
        ("4. 편집", "2–3일", "러프컷, UI 합성, 모션, 음악, 내레이션", "내부 리뷰용 V1 제출"),
        ("5. 검수", "1–2일", "사실성·개인정보·자막·가독성·CTA 검수", "수정 2회 이내 최종 승인"),
        ("6. 배포", "1일", "비율별 렌더, 썸네일, 자막 파일, 업로드", "채널별 링크·UTM·분석 이벤트 확인"),
    ],
    [1250, 1050, 4200, 2860],
    font_size=8.7,
)

add_h2(doc, "제작 레벨 선택")
add_matrix(
    doc,
    ["안", "구성", "장점", "적합한 상황"],
    [
        ("Lean", "내부 촬영 + 화면 녹화 + 템플릿 편집", "빠르고 반복 제작이 쉬움", "시장 반응을 먼저 확인할 때"),
        ("Standard · 추천", "1일 실사 촬영 + 전문 편집 + 실제 UI 모션", "신뢰·완성도·비용의 균형", "공식 런칭과 유료 광고 시작"),
        ("Hero", "감독·촬영팀·배우·세트 + 고급 사운드·다수 로케이션", "강한 브랜드 자산과 PR 활용", "대규모 캠페인·투자/파트너 발표"),
    ],
    [1500, 3200, 2400, 2260],
)
add_text(doc, "예산은 출연, 촬영지, 장비, 내레이션, 음악 라이선스, 수정 횟수에 따라 크게 달라지므로 본 기획 확정 후 2~3개 제작사에 동일 브리프로 비교 견적을 요청한다.",
         size=10, color=MUTED, italic=True, after=8)

# 9. Distribution
add_section_banner(doc, "09", "LAUNCH & DISTRIBUTION")
add_h1(doc, "채널별 역할을 다르게 설계")
add_matrix(
    doc,
    ["채널", "권장 포맷", "콘텐츠 역할", "측정 행동"],
    [
        ("홈페이지", "16:9 60초, 무음 자동재생용 자막본", "브랜드 이해와 제품 확신", "CTA 클릭·가입 시작"),
        ("유튜브", "16:9 60초 + 6초 범퍼", "검색·브랜드 도달", "완주율·사이트 방문"),
        ("인스타그램/쇼츠", "9:16 30초·15초", "3초 훅과 기능 발견", "3초 유지·프로필 클릭"),
        ("사장님 커뮤니티", "9:16 15초 + 설명글", "공감과 입소문", "댓글·저장·상담"),
        ("영업/제휴", "16:9 60초 + 무음 루프", "미팅 오프닝·설명 단축", "데모 요청·파트너 문의"),
    ],
    [1900, 2300, 3000, 2160],
    font_size=8.8,
)

add_h2(doc, "출시 패키지")
add_bullet(doc, "메인 60초: 16:9, 4K 마스터 + 1080p 웹본", bullet_num)
add_bullet(doc, "파생 30초·15초·6초: 9:16, 자막 포함본과 무자막본", bullet_num)
add_bullet(doc, "썸네일 2종: ‘마감 후에도 일하세요?’ / ‘영수증 한 장 찍었을 뿐인데’", bullet_num)
add_bullet(doc, "한글 자막 SRT, 내레이션 없는 무음용 번인 자막, 음악·효과음 분리본", bullet_num)
add_bullet(doc, "CTA QR·단축 URL·UTM 규칙과 채널별 업로드 문구", bullet_num)

add_h2(doc, "성과 지표")
add_matrix(
    doc,
    ["퍼널", "핵심 지표", "초기 판단 기준"],
    [
        ("도달", "3초 시청률·도달 단가", "첫 3초 훅 A/B 중 상대 우위 확인"),
        ("관심", "25%·50%·완주율", "이탈 구간을 찾아 장면 길이 재편집"),
        ("행동", "CTA 클릭률·랜딩 전환", "영상별 UTM으로 가입/상담 기여 추적"),
        ("품질", "댓글·저장·공유·문의 내용", "‘쉽다/필요하다’ 반응과 기능 오해를 함께 수집"),
    ],
    [1500, 3300, 4560],
)

# 10. Approval checklist
add_section_banner(doc, "10", "APPROVAL CHECKLIST")
add_h1(doc, "촬영 전 결정해야 할 것")
for item in [
    "최종 CTA: 무료 체험 / 도입 상담 / 사전 신청 중 하나",
    "대표 타깃 업종: 음식점 중심 또는 카페·소매까지 확장",
    "출연 방식: 실제 고객 / 내부 인원 / 전문 모델",
    "촬영 매장과 날짜, 영업 방해를 피할 시간대",
    "영상에 노출할 실제 제공 기능과 제외 기능",
    "데모 계정·가상 데이터·영수증 샘플 승인",
    "브랜드 슬로건과 로고·앱 아이콘 사용 규칙",
    "내레이션 성별·톤, 음악 라이선스 범위",
    "수정 횟수, 최종 산출물 규격, 원본 프로젝트 보관 기간",
]:
    add_bullet(doc, "□ " + item, bullet_num)

add_h2(doc, "최종 품질 검수")
for item in [
    "첫 3초만 보고도 소상공인 대상 영상임을 알 수 있다.",
    "실제 앱 화면과 내레이션의 기능 설명이 일치한다.",
    "자막을 소리 없이 읽어도 핵심 메시지가 이해된다.",
    "개인정보·실매출·실급여·계좌 정보가 노출되지 않는다.",
    "세무·노무·AI 기능을 과장하거나 결과를 보장하는 표현이 없다.",
    "세로·가로 버전 모두 버튼·자막·QR이 안전영역 안에 있다.",
    "마지막 CTA가 단 하나이며 3초 이상 유지된다.",
    "모바일에서 링크·가입·상담 동선이 실제로 작동한다.",
]:
    add_bullet(doc, "□ " + item, bullet_num)

# Final
add_section_banner(doc, "11", "NEXT ACTION")
add_h1(doc, "바로 실행할 수 있는 다음 단계")
add_callout(
    doc,
    "Go / No-Go Recommendation",
    "‘Standard’ 방식으로 60초 메인과 30초·15초 숏폼을 동시에 제작한다.",
    "첫 캠페인은 영수증 AI 정리 장면을 대표 훅으로 삼고, 손익 대시보드와 직원앱을 신뢰 증거로 연결한다. "
    "촬영 전에 무료 체험 또는 상담 CTA 중 하나를 확정하고, 영상 전용 데모 계정부터 준비한다.",
    fill=PALE_TEAL,
    accent=TEAL,
)
add_h2(doc, "킥오프 아젠다 · 45분")
add_number(doc, "10분 — 목표와 타깃 업종 확정", decimal_num)
add_number(doc, "10분 — 최종 슬로건·CTA 선택", decimal_num)
add_number(doc, "10분 — 노출 기능과 데모 데이터 확정", decimal_num)
add_number(doc, "10분 — 출연·장소·제작 레벨 결정", decimal_num)
add_number(doc, "5분 — 일정·승인자·수정 프로세스 확정", decimal_num)

add_text(doc, "최종 한 문장", size=10, color=BLUE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=7, line=1.0)
add_text(doc, "“복잡한 가게 운영을, 하나로.\n사장님의 숫자 비서, 셈하나.”",
         size=20, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         before=0, after=12, line=1.2)
add_text(doc, "END OF BLUEPRINT", size=8.5, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=0, line=1.0)

doc.core_properties.title = "셈하나 홍보영상 제작 기획서"
doc.core_properties.subject = "60초 브랜드 필름 및 숏폼 파생 콘텐츠 제작 계획"
doc.core_properties.author = "SEMHANA"
doc.core_properties.keywords = "셈하나, 홍보영상, 브랜드필름, 소상공인, 앱마케팅"
doc.core_properties.comments = "narrative_proposal preset with Malgun Gothic Korean readability override"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
